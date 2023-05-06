from docx import Document
import re
# from doc.Section import Section
# from doc.WorkProgram import WorkProgram
import json
import os
# from jinja2 import Template


class WorkProgram:
    def __init__(self):
        self.content = dict()


class Section:
    def __init__(self):
        self.heading = ''
        self.content = ''
        self.links =[]


def read_docx(filename):
    doc = Document(filename)
    content = []

    for p in doc.paragraphs:
        text = p.text.strip()

        if text == '':
            continue
        content.append(text)

    return content


def parse_tables(file):
    doc = Document(file)
    tables_content = []

    for table in doc.tables:
        curr_table = []
        for row in table.rows:
            values = []
            for cell in row.cells:
                for para in cell.paragraphs:
                    values.append(para.text.strip())
            curr_table.append(values)
        tables_content.append(curr_table)

    return tables_content


def parse(doc):
    wp = WorkProgram()
    res = wp.content
    flag = False
    prev = '9.9.9.9.9'

    r = "Титульная страница"
    res[r] = Section()

    for par in doc:

        if par == '' or par == ' ':
            continue

        __match = re.match(r'Раздел [\d]\.', par)  # находим заголовки разделов

        if __match is not None:
            flag = False
            r = __match.group(0)
            res[r] = Section()
            res[r].heading = par[(__match.end() + 1):]
            continue

        _match = re.match(r'\n*[\d]\.[\d]\.[\d]?', par)  # находим заголовки подразделов

        if _match is not None:
            m = _match.group(0)

            res[m] = Section()
            res[m].heading = par[(_match.end() + 1):]

            if not flag:
                res[r].links.append(m)
                prev = m
            elif len(prev) < len(m):
                res[prev].links.append(m)
            else:
                res[r].links.append(m)
                prev = m
            flag = True
        elif flag:
            res[m].content = res[m].content + par + "\n"
        else:
            res[r].content = res[r].content + par + "\n"

    return res


'''''
def access(doc, heading):
    cont = doc.get(heading)
    assert cont is not None
    return cont
'''''


def feed_content(file):
    doc = read_docx(file)
    parsing = parse(doc)

    data = dict()
    for x in parsing:
        data[x] = {
             "title": parsing[x].heading,
             "text": parsing[x].content
        }
        #data[x] = parsing[x].content
        '''''
    tables = parse_tables(file)
    data['table'] = list()
    for x in tables:
        for i in x:
            data['table'].append(i)
'''''
    return data


def feed_structure(file):
    doc = read_docx(file)
    parsing = parse(doc)
    data = dict()
    for x in parsing:
        data[x] = parsing[x].links
    return data


def feed(file):
    # data = feed_structure(file)

    data1 = feed_content(file)

    # with open("results_structure\\" + file[28:-5] + "_structure.json", "w", encoding='utf-8') as write_file:
    #     json.dump(data, write_file, indent=4, ensure_ascii=False)
    with open("results_content2\\" + file[7:-5] + "_content.json", "w", encoding='utf-8') as write_file:
        json.dump(data1, write_file, indent=4, ensure_ascii=False)

'''''
def template(file):
    doc = read_docx(file)
    parsing = parse(doc)
    tem = open('template.txt').read()
    template = Template(tem)
    mydoc = Document()
    mydoc.add_paragraph(template.render(data=parsing))
    mydoc.save("result\\" + file[5:-5] + "_res.docx")
'''''

if __name__ == '__main__':

    for f in os.listdir('2 data/'):
        print(f)
        feed('2 data/' + f)




